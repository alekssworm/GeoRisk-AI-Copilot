from __future__ import annotations

import logging
from datetime import datetime, timezone
from io import BytesIO
from xml.sax.saxutils import escape

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5F6B7A"
RISK_RED = "9B1C1C"
logger = logging.getLogger("georisk.reporting")


def _plain_text(value: object) -> str:
    return str(value or "").replace("**", "").replace("`", "").strip()


def _report_timestamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _summary_rows(report: dict) -> list[tuple[str, str]]:
    prediction = report["prediction"]
    uncertainty = prediction.get("uncertainty") or {}
    interval = "Not available"
    if uncertainty.get("lower_usv_h") is not None:
        interval = (
            f"{uncertainty['lower_usv_h']:.3f}-{uncertainty['upper_usv_h']:.3f} uSv/h "
            f"({uncertainty['confidence_label']})"
        )
    return [
        ("Predicted dose rate", f"{prediction['dose_rate_usv_h']:.3f} uSv/h"),
        ("Risk level", prediction["risk_level"]),
        ("Model spread (P10-P90)", interval),
        ("Model version", prediction["model_version"]),
    ]


def build_risk_report_pdf(report: dict) -> bytes:
    from matplotlib import font_manager
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        KeepTogether,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    regular_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    try:
        regular_path = font_manager.findfont("DejaVu Sans")
        bold_path = font_manager.findfont(
            font_manager.FontProperties(family="DejaVu Sans", weight="bold")
        )
        if "GeoRiskSans" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("GeoRiskSans", regular_path))
            pdfmetrics.registerFont(TTFont("GeoRiskSans-Bold", bold_path))
        regular_font = "GeoRiskSans"
        bold_font = "GeoRiskSans-Bold"
    except (OSError, RuntimeError, ValueError) as exc:
        logger.debug("DejaVu font registration failed; using PDF core fonts: %s", exc)

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.8 * inch,
        bottomMargin=0.75 * inch,
        title="GeoRisk AI Copilot Risk Analysis",
        author="GeoRisk AI Copilot",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "GeoRiskBody",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
        textColor=colors.HexColor("#1F2933"),
    )
    title = ParagraphStyle(
        "GeoRiskTitle",
        parent=body,
        fontName=bold_font,
        fontSize=23,
        leading=27,
        textColor=colors.HexColor(f"#{NAVY}"),
        spaceAfter=4,
    )
    subtitle = ParagraphStyle(
        "GeoRiskSubtitle",
        parent=body,
        fontSize=11,
        textColor=colors.HexColor(f"#{MUTED}"),
        spaceAfter=16,
    )
    heading = ParagraphStyle(
        "GeoRiskHeading",
        parent=body,
        fontName=bold_font,
        fontSize=15,
        leading=18,
        textColor=colors.HexColor(f"#{BLUE}"),
        spaceBefore=12,
        spaceAfter=7,
    )
    small = ParagraphStyle(
        "GeoRiskSmall",
        parent=body,
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor(f"#{MUTED}"),
    )
    center_metric = ParagraphStyle(
        "GeoRiskMetric",
        parent=body,
        alignment=TA_CENTER,
        fontName=bold_font,
        fontSize=11,
        leading=14,
    )
    left_table = ParagraphStyle(
        "GeoRiskTable",
        parent=small,
        alignment=TA_LEFT,
        fontSize=8.5,
        leading=10.5,
    )

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#D8DEE8"))
        canvas.setLineWidth(0.5)
        canvas.line(inch, letter[1] - 0.55 * inch, letter[0] - inch, letter[1] - 0.55 * inch)
        canvas.setFont(regular_font, 8)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(inch, letter[1] - 0.42 * inch, "GeoRisk AI Copilot | Risk Analysis")
        canvas.drawRightString(letter[0] - inch, 0.42 * inch, f"Page {doc.page}")
        canvas.restoreState()

    prediction = report["prediction"]
    story = [
        Spacer(1, 0.1 * inch),
        Paragraph("RISK ANALYSIS", title),
        Paragraph(f"Environmental radiation screening report | {_report_timestamp()}", subtitle),
    ]

    summary_data = [
        [
            Paragraph("DOSE RATE", small),
            Paragraph("RISK LEVEL", small),
            Paragraph("MODEL CONFIDENCE", small),
        ],
        [
            Paragraph(f"{prediction['dose_rate_usv_h']:.3f} uSv/h", center_metric),
            Paragraph(escape(prediction["risk_level"]), center_metric),
            Paragraph(
                escape((prediction.get("uncertainty") or {}).get("confidence_label", "n/a")),
                center_metric,
            ),
        ],
    ]
    summary = Table(summary_data, colWidths=[2.15 * inch] * 3, repeatRows=1)
    summary.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_BLUE}")),
                ("BACKGROUND", (0, 1), (-1, 1), colors.white),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#C7D2E3")),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D8DEE8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend([summary, Spacer(1, 0.12 * inch)])

    story.append(Paragraph("Executive Summary", heading))
    story.append(
        Paragraph(
            escape(
                f"The modeled dose rate is {prediction['dose_rate_usv_h']:.3f} uSv/h, "
                f"classified as {prediction['risk_level']}. {prediction['advisory']}"
            ),
            body,
        )
    )
    uncertainty = prediction.get("uncertainty") or {}
    if uncertainty.get("lower_usv_h") is not None:
        story.append(
            Paragraph(
                escape(
                    f"Tree-ensemble P10-P90 spread: {uncertainty['lower_usv_h']:.3f}-"
                    f"{uncertainty['upper_usv_h']:.3f} uSv/h. This is an uncalibrated "
                    "model-dispersion indicator, not a guaranteed confidence interval."
                ),
                body,
            )
        )
    distribution = prediction.get("distribution_check") or {}
    if distribution.get("warning"):
        warning_table = Table(
            [[Paragraph(f"<b>Extrapolation warning:</b> {escape(distribution['warning'])}", body)]],
            colWidths=[6.5 * inch],
        )
        warning_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF4E5")),
                    ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#D89B2B")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.extend([warning_table, Spacer(1, 0.08 * inch)])

    story.append(Paragraph("Scenario Comparison", heading))
    scenario_rows = [["Scenario", "Dose (uSv/h)", "Risk", "Delta"]]
    for row in report.get("scenario_comparison", []):
        scenario_rows.append(
            [
                _plain_text(row.get("name")),
                f"{float(row.get('dose_rate_usv_h', 0)):.3f}",
                _plain_text(row.get("risk_level")),
                f"{float(row.get('delta_vs_baseline_usv_h', 0)):+.3f}",
            ]
        )
    scenario_table = Table(
        [[Paragraph(escape(str(cell)), left_table) for cell in row] for row in scenario_rows],
        colWidths=[2.55 * inch, 1.25 * inch, 1.25 * inch, 1.35 * inch],
        repeatRows=1,
    )
    scenario_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_GRAY}")),
                ("FONTNAME", (0, 0), (-1, 0), bold_font),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7CDD6")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(scenario_table)

    story.append(Paragraph("Main Risk Drivers", heading))
    drivers = report.get("explanation", {}).get("top_features", [])[:6]
    driver_items = [
        ListItem(
            Paragraph(
                escape(
                    f"{item['feature']}: {item['direction']} (value {float(item['value']):.3f})"
                ),
                body,
            )
        )
        for item in drivers
    ]
    if driver_items:
        story.append(ListFlowable(driver_items, bulletType="bullet", leftIndent=18))
    else:
        story.append(Paragraph("No feature explanation is available.", body))

    story.append(Paragraph("Document-Grounded Context", heading))
    rag_answer = report.get("rag_answer")
    if rag_answer:
        answer_text = (
            rag_answer.get("answer") if isinstance(rag_answer, dict) else rag_answer.answer
        )
        story.append(Paragraph(escape(_plain_text(answer_text)), body))
        citations = (
            rag_answer.get("citations", [])
            if isinstance(rag_answer, dict)
            else rag_answer.citations
        )
        for citation in citations:
            story.append(Paragraph(escape(_plain_text(citation.get("label"))), small))
    else:
        story.append(Paragraph("No document question was included.", body))

    story.append(Paragraph("Recommended Next Actions", heading))
    actions = [
        "Validate contamination inputs against field survey measurements.",
        "Repeat the comparison with conservative rainfall and runoff assumptions.",
        "Confirm intervention thresholds and monitoring cadence with qualified experts.",
    ]
    story.append(
        ListFlowable(
            [ListItem(Paragraph(escape(action), body)) for action in actions],
            bulletType="bullet",
            leftIndent=18,
        )
    )

    disclaimer = Table(
        [
            [
                Paragraph(
                    "<b>Decision-use limitation.</b> This report is a screening aid. It does "
                    "not replace calibrated instruments, laboratory analysis, site-specific "
                    "dose assessment, or advice from radiation-protection professionals.",
                    small,
                )
            ]
        ],
        colWidths=[6.5 * inch],
    )
    disclaimer.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F8FA")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#C7CDD6")),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend([Spacer(1, 0.12 * inch), KeepTogether(disclaimer)])
    document.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
    return buffer.getvalue()


def build_risk_report_docx(report: dict) -> bytes:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def set_run(run, size=11, color="1F2933", bold=False, italic=False):
        run.font.name = "Calibri"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold
        run.italic = italic

    def shade_cell(cell, fill):
        properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)

    def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
        properties = cell._tc.get_or_add_tcPr()
        margins = properties.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            properties.append(margins)
        for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
            node = margins.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def set_table_widths(table, widths):
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
                set_cell_margins(row.cells[index])
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("GeoRisk AI Copilot | Risk Analysis"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Generated by GeoRisk AI Copilot | Page "), size=8, color=MUTED)
    page_run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    page_run._r.extend([begin, instruction, end])

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(4)
    set_run(title_paragraph.add_run("RISK ANALYSIS"), size=23, color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run(
        subtitle.add_run(f"Environmental radiation screening report | {_report_timestamp()}"),
        size=11,
        color=MUTED,
    )

    prediction = report["prediction"]
    summary = document.add_table(rows=2, cols=3)
    summary.style = "Table Grid"
    set_table_widths(summary, [2.16, 2.17, 2.17])
    labels = ("DOSE RATE", "RISK LEVEL", "MODEL CONFIDENCE")
    values = (
        f"{prediction['dose_rate_usv_h']:.3f} uSv/h",
        prediction["risk_level"],
        (prediction.get("uncertainty") or {}).get("confidence_label", "n/a"),
    )
    for index, label in enumerate(labels):
        shade_cell(summary.cell(0, index), LIGHT_BLUE)
        paragraph = summary.cell(0, index).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(paragraph.add_run(label), size=8.5, color=MUTED, bold=True)
        value_paragraph = summary.cell(1, index).paragraphs[0]
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(value_paragraph.add_run(str(values[index])), size=11, color=NAVY, bold=True)

    document.add_heading("Executive Summary", level=1)
    paragraph = document.add_paragraph()
    set_run(
        paragraph.add_run(
            f"The modeled dose rate is {prediction['dose_rate_usv_h']:.3f} uSv/h, "
            f"classified as {prediction['risk_level']}. {prediction['advisory']}"
        )
    )
    uncertainty = prediction.get("uncertainty") or {}
    if uncertainty.get("lower_usv_h") is not None:
        paragraph = document.add_paragraph()
        set_run(
            paragraph.add_run(
                f"Tree-ensemble P10-P90 spread: {uncertainty['lower_usv_h']:.3f}-"
                f"{uncertainty['upper_usv_h']:.3f} uSv/h. This is an uncalibrated "
                "model-dispersion indicator, not a guaranteed confidence interval."
            )
        )
    distribution = prediction.get("distribution_check") or {}
    if distribution.get("warning"):
        warning = document.add_table(rows=1, cols=1)
        warning.style = "Table Grid"
        set_table_widths(warning, [6.5])
        shade_cell(warning.cell(0, 0), "FFF4E5")
        paragraph = warning.cell(0, 0).paragraphs[0]
        set_run(paragraph.add_run("Extrapolation warning: "), bold=True, color=RISK_RED)
        set_run(paragraph.add_run(distribution["warning"]))

    document.add_heading("Scenario Comparison", level=1)
    rows = report.get("scenario_comparison", [])
    scenario_table = document.add_table(rows=1, cols=4)
    scenario_table.style = "Table Grid"
    set_table_widths(scenario_table, [2.5, 1.25, 1.25, 1.5])
    for index, label in enumerate(("Scenario", "Dose (uSv/h)", "Risk", "Delta")):
        shade_cell(scenario_table.cell(0, index), LIGHT_GRAY)
        set_run(scenario_table.cell(0, index).paragraphs[0].add_run(label), size=9, bold=True)
    for row in rows:
        cells = scenario_table.add_row().cells
        values = (
            _plain_text(row.get("name")),
            f"{float(row.get('dose_rate_usv_h', 0)):.3f}",
            _plain_text(row.get("risk_level")),
            f"{float(row.get('delta_vs_baseline_usv_h', 0)):+.3f}",
        )
        for index, value in enumerate(values):
            set_cell_margins(cells[index])
            set_run(cells[index].paragraphs[0].add_run(value), size=9)

    document.add_heading("Main Risk Drivers", level=1)
    drivers = report.get("explanation", {}).get("top_features", [])[:6]
    if drivers:
        list_style = document.styles["List Bullet"]
        list_style.paragraph_format.left_indent = Inches(0.5)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)
        list_style.paragraph_format.space_after = Pt(8)
        list_style.paragraph_format.line_spacing = 1.167
        for item in drivers:
            paragraph = document.add_paragraph(style="List Bullet")
            set_run(
                paragraph.add_run(
                    f"{item['feature']}: {item['direction']} (value {float(item['value']):.3f})"
                )
            )
    else:
        document.add_paragraph("No feature explanation is available.")

    document.add_heading("Document-Grounded Context", level=1)
    rag_answer = report.get("rag_answer")
    if rag_answer:
        answer_text = (
            rag_answer.get("answer") if isinstance(rag_answer, dict) else rag_answer.answer
        )
        document.add_paragraph(_plain_text(answer_text))
        citations = (
            rag_answer.get("citations", [])
            if isinstance(rag_answer, dict)
            else rag_answer.citations
        )
        for citation in citations:
            paragraph = document.add_paragraph()
            set_run(paragraph.add_run(_plain_text(citation.get("label"))), size=9, color=MUTED)
    else:
        document.add_paragraph("No document question was included.")

    document.add_heading("Recommended Next Actions", level=1)
    for action in (
        "Validate contamination inputs against field survey measurements.",
        "Repeat the comparison with conservative rainfall and runoff assumptions.",
        "Confirm intervention thresholds and monitoring cadence with qualified experts.",
    ):
        paragraph = document.add_paragraph(style="List Bullet")
        set_run(paragraph.add_run(action))

    disclaimer = document.add_table(rows=1, cols=1)
    disclaimer.style = "Table Grid"
    set_table_widths(disclaimer, [6.5])
    shade_cell(disclaimer.cell(0, 0), "F7F8FA")
    paragraph = disclaimer.cell(0, 0).paragraphs[0]
    set_run(paragraph.add_run("Decision-use limitation. "), size=9, bold=True)
    set_run(
        paragraph.add_run(
            "This report is a screening aid. It does not replace calibrated instruments, "
            "laboratory analysis, site-specific dose assessment, or advice from "
            "radiation-protection professionals."
        ),
        size=9,
        color=MUTED,
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
